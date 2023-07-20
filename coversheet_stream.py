import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, RGBColor
import streamlit as st
from io import BytesIO

st.set_page_config(page_title = "Coversheet Document", layout = "wide")

st.title("Coversheet Generator")
st.text(""" Please provide the **Originator's Document ID** of the document you want to issue in the first field. 
In case you also want to place **Reference Documents** in the coveersheet please specify how many you wish
to place in the second field and provide one **Originator's Document ID** per field. 

When you are ready click on the **Generate Coversheet** button. After this please click the **Download Coversheet** 
button. You will find the Coversheet of the document in question in your **Downloads** folder.

If you have any questions please contact c.topalidis@iv-one.nl""")

# Read the data from the Excel file
path_mdr = "https://raw.githubusercontent.com/christos1991t/python_test/master/" \
           "Copy%20of%20NATIVE_IJ2-COM-000-GSC-H1001-0011-LD_02-Phase%20Prep%20MDR%20Setup%20%20Concept%20Design.xlsx"

df = pd.read_excel(path_mdr, sheet_name = "MDR", engine = 'openpyxl')

path_cover = "Coversheet.docx"


# Load the Word template
doc = Document(path_cover)

# Get the table and cell to populate
table = doc.tables[0]
cell_to_populate1 = table.cell(0, 1)
cell_to_populate2 = table.cell(1, 0)
cell_to_populate3 = table.cell(1, 1)

cell_to_populate4 = table.cell(2, 0)
cell_to_populate5 = table.cell(2, 1)
cell_to_populate6 = table.cell(2, 2)

cell_to_populate7 = table.cell(3, 0)
cell_to_populate8 = table.cell(3, 2)

cell_to_populate9 = table.cell(4, 0)
cell_to_populate10 = table.cell(4, 2)

cell_to_populate11 = table.cell(5, 0)
cell_to_populate12 = table.cell(5, 2)

cell_to_populate13 = table.cell(6, 0)
cell_to_populate14 = table.cell(6, 2)

cell_to_populate15 = table.cell(7, 0)
cell_to_populate16 = table.cell(7, 2)

cell_to_populate17 = table.cell(8, 0)
cell_to_populate18 = table.cell(8, 2)

tennet = table.cell(9, 0)
consortium = table.cell(9, 2)

project_title = table.cell(1, 2)

header = doc.sections[1].header
table_1 = header.tables[0]

header_cell = table_1.cell(1, 1)

table_page_2 = doc.tables[3]


user = st.text_input("Please provide Originator's Document ID ", key = "first_input")
trig = str(user).strip()


def get_cell_value(df, index , name):
    value = df.loc[index, name]
    return "N/A" if pd.isna(value) else str(value)


export_title = ""


for i, row in df.iterrows():
    if trig == row["Originator's Document ID"]:

        cell_value1 = get_cell_value(df, i, "Document Title")
        cell_value2 = get_cell_value(df, i, "TenneT Document ID")
        cell_value3 = get_cell_value(df, i, "TenneT Revision")

        cell_value4 = get_cell_value(df, i, "Originator's Document ID")
        cell_value5 = get_cell_value(df, i, "Originator's Revision")
        cell_value6 = get_cell_value(df, i, "Asset Document Reference")

        cell_value9 = get_cell_value(df, i, "Purpose of Submission")
        cell_value10 = get_cell_value(df, i, "WBS Code")

        cell_value11 = get_cell_value(df, i, "Purpose of Issue")
        cell_value12 = get_cell_value(df, i, "WBS Name")

        cell_value13 = get_cell_value(df, i, "Book")
        cell_value14 = get_cell_value(df, i, "DCC#")

        cell_value15 = get_cell_value(df, i, "Chapter")
        cell_value16 = get_cell_value(df, i, "Document Kind")

        cell_value17 = get_cell_value(df, i, "Subchapter")
        cell_value18 = get_cell_value(df, i, "Security Level")

        tennet_beeld = "Tennet.jpg"
        consortium_beeld = "Consortium.png"

        # Populate the cell in the Word document

        cell_to_populate1.text = "Document Title - "
        cell_to_populate1.add_paragraph().add_run(str(cell_value1)).font.color.rgb = RGBColor(0, 112,
                                                                                              192)  # Light blue color
        cell_to_populate2.text = str(f"TenneT Document ID - ")
        cell_to_populate2.add_paragraph().add_run(str(cell_value2)).font.color.rgb = RGBColor(0, 112,
                                                                                              192)  # Light blue color
        cell_to_populate3.text = str(f"TenneT revision - ")
        cell_to_populate3.add_paragraph().add_run(str(cell_value3)).font.color.rgb = RGBColor(0, 112,
                                                                                              192)  # Light blue color

        cell_to_populate4.text = str(f"Originator´s Document ID - ")
        cell_to_populate4.add_paragraph().add_run(str(cell_value4)).font.color.rgb = RGBColor(0, 112,
                                                                                              192)  # Light blue color
        cell_to_populate5.text = str(f"Originator´s Revision - ")
        cell_to_populate5.add_paragraph().add_run(str(cell_value5)).font.color.rgb = RGBColor(0, 112,
                                                                                              192)  # Light blue color
        cell_to_populate6.text = str(f"Asset Document Reference - ")
        cell_to_populate6.add_paragraph().add_run(str(cell_value6)).font.color.rgb = RGBColor(0, 112,
                                                                                              192)  # Light blue color

        cell_to_populate7.text = str(f"Contractor - ")
        cell_to_populate7.add_paragraph().add_run(str("GSC")).font.color.rgb = RGBColor(0, 112,
                                                                                              192)  # Light blue color
        cell_to_populate8.text = str(f"Contract Number - ")
        cell_to_populate8.add_paragraph().add_run(str("GSC-2GW")).font.color.rgb = RGBColor(0, 112,
                                                                                      192)  # Light blue color

        cell_to_populate9.text = str(f"Purpose of Submission - ")
        cell_to_populate9.add_paragraph().add_run(str(cell_value9)).font.color.rgb = RGBColor(0, 112,
                                                                                            192)  # Light blue color
        cell_to_populate10.text = str(f"WBS Code - ")
        cell_to_populate10.add_paragraph().add_run(str(cell_value10)).font.color.rgb = RGBColor(0, 112,
                                                                                              192)  # Light blue color

        cell_to_populate11.text = str(f"Purpose of Issue - ")
        cell_to_populate11.add_paragraph().add_run(str(cell_value11)).font.color.rgb = RGBColor(0, 112,
                                                                                                192)  # Light blue color
        cell_to_populate12.text = str(f"WBS Name - ")
        cell_to_populate12.add_paragraph().add_run(str(cell_value12)).font.color.rgb = RGBColor(0, 112,
                                                                                                192)  # Light blue color

        cell_to_populate13.text = str(f"Book - ")
        cell_to_populate13.add_paragraph().add_run(str(cell_value13)).font.color.rgb = RGBColor(0, 112,
                                                                                                192)  # Light blue color
        cell_to_populate14.text = str(f"DCC# - ")
        cell_to_populate14.add_paragraph().add_run(str(cell_value14)).font.color.rgb = RGBColor(0, 112,
                                                                                                192)  # Light blue color

        cell_to_populate15.text = str(f"Chapter - ")
        cell_to_populate15.add_paragraph().add_run(str(cell_value15)).font.color.rgb = RGBColor(0, 112,
                                                                                                192)  # Light blue color
        cell_to_populate16.text = str(f"Chapter - ")
        cell_to_populate16.add_paragraph().add_run(str(cell_value16)).font.color.rgb = RGBColor(0, 112,
                                                                                                192)  # Light blue color

        cell_to_populate17.text = str(f"Subchapter - ")
        cell_to_populate17.add_paragraph().add_run(str(cell_value17)).font.color.rgb = RGBColor(0, 112,
                                                                                                192)  # Light blue color
        cell_to_populate18.text = str(f"Security Level - ")
        cell_to_populate18.add_paragraph().add_run(str(cell_value18)).font.color.rgb = RGBColor(0, 112,
                                                                                                192)  # Light blue color

        tennet_paragraph = tennet.paragraphs[0]
        tennet_run = tennet_paragraph.add_run()
        tennet_run.add_picture(tennet_beeld)
        tennet.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        consortium_paragraph = consortium.paragraphs[0]
        consortium_run = consortium_paragraph.add_run()
        consortium_run.add_picture(consortium_beeld)
        consortium.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        header_cell.text = str(f"TenneT Document ID - \n{cell_value2}")
        project_title.text = str(f"Project Name - ")
        project_title.add_paragraph().add_run(str("IJmuiden Ver Beta")).font.color.rgb = RGBColor(0, 112,
                                                                                            192)  # Light blue color

        export_title = cell_value1

    else:
        pass

last_filled_row = 1
counter = 1

# Prompt the user to specify how many unique codes to insert
ref_doc_num = st.number_input("How many referenced documents do you wish to add?", value=1)

# List to store the user-inputted codes
user_codes = []

# Loop to get user input for unique codes
for i in range(ref_doc_num):
    # Generate a unique key for each text_input widget
    user_code = st.text_input(f"Please provide the Originator's Document ID ({i+1}):")
    user_codes.append(user_code.strip())

print(user_codes)

# Process the user codes and populate the table if found in the Excel file
for user_code in user_codes:
    for j, row in df.iterrows():
        if user_code == row["Originator's Document ID"]:
            cell_value_ref_1 = df.loc[j, "TenneT Document ID"]
            cell_value_ref_2 = df.loc[j, "Originator's Document ID"]
            cell_value_ref_3 = df.loc[j, "Document Title"]

            # Find the next empty row to insert data after the last filled row
            for k, table_page_2_row in enumerate(table_page_2.rows[last_filled_row + 1:], start=last_filled_row + 1):
                is_row_filled = all(cell.text.strip() for cell in table_page_2_row.cells)
                if not is_row_filled:
                    cell_table_2_1_counter = table_page_2_row.cells[0]
                    cell_table_2_1_counter.text = str(counter)

                    cell_table_2_1 = table_page_2_row.cells[1]
                    cell_table_2_1.text = cell_value_ref_1
                    cell_table_2_2 = table_page_2_row.cells[2]
                    cell_table_2_2.text = cell_value_ref_2
                    cell_table_2_3 = table_page_2_row.cells[3]
                    cell_table_2_3.text = cell_value_ref_3

                    # Update last_filled_row to the current row index
                    last_filled_row = k
                    counter += 1
                    break  # Exit the loop after inserting the reference
            else:
                # If the loop completes without finding an empty row, add a new row
                new_row = table_page_2.add_row()
                cell_table_2_1_counter = new_row.cells[0]
                cell_table_2_1_counter.text = str(counter)
                cell_table_2_1 = new_row.cells[1]
                cell_table_2_1.text = cell_value_ref_1
                cell_table_2_2 = new_row.cells[2]
                cell_table_2_2.text = cell_value_ref_2
                cell_table_2_3 = new_row.cells[3]
                cell_table_2_3.text = cell_value_ref_3

                # Update last_filled_row to the current row index
                last_filled_row = len(table_page_2.rows) - 1
                counter += 1


#dummy button
if st.button("Generate Coversheet"):

    # Save the Word document to a BytesIO stream
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    # Provide the Word document as a downloadable file
    st.download_button(
        label="Download Coversheet",
        data=doc_stream,
        file_name=f"Coversheet - {str(export_title)}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )